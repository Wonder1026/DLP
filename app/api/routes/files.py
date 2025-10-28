from fastapi import APIRouter, UploadFile, File, Depends, HTTPException, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pathlib import Path
import uuid
from datetime import datetime
from app.database import get_db
from app.models.file import UploadedFile
from app.models.user import User
from app.websocket.manager import manager
from fastapi.responses import FileResponse as FastAPIFileResponse
from docx import Document


router = APIRouter()

# Папка для загрузки файлов
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Разрешённые типы файлов
ALLOWED_EXTENSIONS = {".exe", ".doc", ".docx"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


@router.post("/upload")
async def upload_file(
        user_id: int,
        moderation_type: str = "manual",  # ← добавили параметр
        file: UploadFile = File(...),
        db: AsyncSession = Depends(get_db)
):
    """Загрузка файла"""

    # Проверяем пользователя
    result = await db.execute(select(User).where(User.id == user_id))
    user = result.scalar_one_or_none()

    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Пользователь не найден"
        )

    # Проверяем тип модерации
    if moderation_type not in ["manual", "virustotal"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Недопустимый тип модерации. Разрешены: manual, virustotal"
        )

    # Проверяем расширение файла
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Недопустимый тип файла. Разрешены: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # Читаем файл
    content = await file.read()
    file_size = len(content)

    # Проверяем размер
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Файл слишком большой. Максимум: {MAX_FILE_SIZE // 1024 // 1024} MB"
        )

    # Генерируем уникальное имя файла
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    file_path = UPLOAD_DIR / unique_filename

    # Сохраняем файл
    with open(file_path, "wb") as f:
        f.write(content)

    # Сохраняем информацию в БД
    uploaded_file = UploadedFile(
        user_id=user.id,
        username=user.username,
        display_name=user.display_name,
        filename=file.filename,
        file_path=str(file_path),
        file_size=file_size,
        file_type=file_ext.replace(".", ""),
        mime_type=file.content_type,
        status="pending",
        moderation_type=moderation_type  # ← сохраняем тип модерации
    )

    db.add(uploaded_file)
    await db.commit()
    await db.refresh(uploaded_file)

    moderation_text = "ручную модерацию" if moderation_type == "manual" else "проверку VirusTotal"
    print(f"📎 Файл загружен: {file.filename} от {user.display_name} (тип модерации: {moderation_type})")

    return {
        "status": "success",
        "message": f"Файл загружен и отправлен на {moderation_text}",
        "file": uploaded_file.to_dict()
    }


@router.get("/list")
async def get_files(user_id: int, db: AsyncSession = Depends(get_db)):
    """Получить список файлов пользователя"""

    result = await db.execute(
        select(UploadedFile)
        .where(UploadedFile.user_id == user_id)
        .order_by(UploadedFile.created_at.desc())
    )
    files = result.scalars().all()

    return {
        "files": [f.to_dict() for f in files]
    }


@router.get("/pending")
async def get_pending_files(admin_id: int, db: AsyncSession = Depends(get_db)):
    """Получить файлы на модерации (только для админов)"""

    # Проверяем права админа
    result = await db.execute(select(User).where(User.id == admin_id))
    admin = result.scalar_one_or_none()

    if not admin or not admin.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещён. Только для администраторов."
        )

    # Получаем файлы на модерации
    result = await db.execute(
        select(UploadedFile)
        .where(UploadedFile.status == "pending")
        .order_by(UploadedFile.created_at.desc())
    )
    files = result.scalars().all()

    return {
        "files": [f.to_dict() for f in files],
        "count": len(files)
    }


@router.get("/all")
async def get_all_files(admin_id: int, db: AsyncSession = Depends(get_db)):
    """Получить все файлы (только для админов)"""

    # Проверяем права админа
    result = await db.execute(select(User).where(User.id == admin_id))
    admin = result.scalar_one_or_none()

    if not admin or not admin.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещён. Только для администраторов."
        )

    # Получаем все файлы
    result = await db.execute(
        select(UploadedFile).order_by(UploadedFile.created_at.desc())
    )
    files = result.scalars().all()

    return {
        "files": [f.to_dict() for f in files],
        "count": len(files)
    }


@router.post("/{file_id}/approve")
async def approve_file(
        file_id: int,
        admin_id: int,
        db: AsyncSession = Depends(get_db)
):
    """Одобрить файл"""

    # Проверяем права админа
    result = await db.execute(select(User).where(User.id == admin_id))
    admin = result.scalar_one_or_none()

    if not admin or not admin.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещён"
        )

    # Находим файл
    result = await db.execute(select(UploadedFile).where(UploadedFile.id == file_id))
    file_obj = result.scalar_one_or_none()

    if not file_obj:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )

    file_obj.status = "approved"
    await db.commit()
    await db.refresh(file_obj)

    print(f"✅ Файл одобрен: {file_obj.filename}")

    # 📢 Отправляем уведомление всем через WebSocket
    await manager.broadcast({
        "type": "file_status_update",
        "file_id": file_obj.id,
        "status": "approved"
    })

    return {
        "status": "success",
        "message": f"Файл '{file_obj.filename}' одобрен",
        "file": file_obj.to_dict()
    }


@router.get("/pending")
async def get_pending_files(admin_id: int, db: AsyncSession = Depends(get_db)):
    """Получить файлы на модерации (только для админов)"""

    # Проверяем права админа
    result = await db.execute(select(User).where(User.id == admin_id))
    admin = result.scalar_one_or_none()

    if not admin or not admin.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещён. Только для администраторов."
        )

    # Получаем файлы на модерации
    result = await db.execute(
        select(UploadedFile)
        .where(UploadedFile.status == "pending")
        .order_by(UploadedFile.created_at.desc())
    )
    files = result.scalars().all()

    return {
        "files": [f.to_dict() for f in files],
        "count": len(files)
    }


@router.get("/all")
async def get_all_files(admin_id: int, db: AsyncSession = Depends(get_db)):
    """Получить все файлы (только для админов)"""

    # Проверяем права админа
    result = await db.execute(select(User).where(User.id == admin_id))
    admin = result.scalar_one_or_none()

    if not admin or not admin.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещён. Только для администраторов."
        )

    # Получаем все файлы
    result = await db.execute(
        select(UploadedFile).order_by(UploadedFile.created_at.desc())
    )
    files = result.scalars().all()

    return {
        "files": [f.to_dict() for f in files],
        "count": len(files)
    }


@router.post("/{file_id}/approve")
async def approve_file(
        file_id: int,
        admin_id: int,
        db: AsyncSession = Depends(get_db)
):
    """Одобрить файл"""

    # Проверяем права админа
    result = await db.execute(select(User).where(User.id == admin_id))
    admin = result.scalar_one_or_none()

    if not admin or not admin.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещён"
        )

    # Находим файл
    result = await db.execute(select(UploadedFile).where(UploadedFile.id == file_id))
    file_obj = result.scalar_one_or_none()

    if not file_obj:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )

    file_obj.status = "approved"
    await db.commit()

    print(f"✅ Файл одобрен: {file_obj.filename}")

    return {
        "status": "success",
        "message": f"Файл '{file_obj.filename}' одобрен",
        "file": file_obj.to_dict()
    }


@router.post("/{file_id}/reject")
async def reject_file(
        file_id: int,
        admin_id: int,
        db: AsyncSession = Depends(get_db)
):
    """Отклонить файл"""

    # Проверяем права админа
    result = await db.execute(select(User).where(User.id == admin_id))
    admin = result.scalar_one_or_none()

    if not admin or not admin.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещён"
        )

    # Находим файл
    result = await db.execute(select(UploadedFile).where(UploadedFile.id == file_id))
    file_obj = result.scalar_one_or_none()

    if not file_obj:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )

    file_obj.status = "rejected"
    await db.commit()
    await db.refresh(file_obj)

    print(f"❌ Файл отклонён: {file_obj.filename}")

    # 📢 Отправляем уведомление всем через WebSocket
    await manager.broadcast({
        "type": "file_status_update",
        "file_id": file_obj.id,
        "status": "rejected"
    })

    return {
        "status": "success",
        "message": f"Файл '{file_obj.filename}' отклонён",
        "file": file_obj.to_dict()
    }


@router.get("/approved")
async def get_approved_files(db: AsyncSession = Depends(get_db)):
    """Получить все одобренные файлы"""

    result = await db.execute(
        select(UploadedFile)
        .where(UploadedFile.status == "approved")
        .order_by(UploadedFile.created_at)
    )
    files = result.scalars().all()

    return {
        "files": [f.to_dict() for f in files]
    }


@router.get("/my-files")
async def get_my_files(user_id: int, db: AsyncSession = Depends(get_db)):
    """Получить файлы текущего пользователя (включая pending)"""

    result = await db.execute(
        select(UploadedFile)
        .where(UploadedFile.user_id == user_id)
        .order_by(UploadedFile.created_at)
    )
    files = result.scalars().all()

    return {
        "files": [f.to_dict() for f in files]
    }


@router.post("/{file_id}/check-virustotal")
async def check_virustotal(
        file_id: int,
        admin_id: int,
        db: AsyncSession = Depends(get_db)
):
    """Проверить файл через VirusTotal API"""

    # Проверяем права админа
    result = await db.execute(select(User).where(User.id == admin_id))
    admin = result.scalar_one_or_none()

    if not admin or not admin.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещён"
        )

    # Находим файл
    result = await db.execute(select(UploadedFile).where(UploadedFile.id == file_id))
    file_obj = result.scalar_one_or_none()

    if not file_obj:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )

    # Проверяем через VirusTotal
    from app.services.virustotal_service import virustotal_service
    import json

    result = await virustotal_service.scan_file(file_obj.file_path)

    # Сохраняем результат
    file_obj.virustotal_result = json.dumps(result, ensure_ascii=False)

    # Автоматически одобряем/отклоняем в зависимости от результата
    if result.get("status") == "clean":
        file_obj.status = "approved"
        print(f"✅ Файл автоматически одобрен (VirusTotal: чисто)")
    elif result.get("status") == "malicious":
        file_obj.status = "rejected"
        print(f"❌ Файл автоматически отклонён (VirusTotal: обнаружены вирусы)")
    else:
        # Если подозрительный или ошибка - оставляем на ручную проверку
        print(f"⚠️ Файл требует ручной проверки (VirusTotal: {result.get('status')})")

    await db.commit()
    await db.refresh(file_obj)

    # Отправляем обновление статуса через WebSocket
    from app.websocket.manager import manager
    await manager.broadcast({
        "type": "file_status_update",
        "file_id": file_obj.id,
        "status": file_obj.status
    })

    return {
        "status": "success",
        "message": "Проверка VirusTotal завершена",
        "virustotal_result": result,
        "file": file_obj.to_dict()
    }





@router.get("/{file_id}/download")
async def download_file(
        file_id: int,
        user_id: int,
        db: AsyncSession = Depends(get_db)
):
    """Скачать файл (только одобренные)"""

    # Проверяем пользователя
    result = await db.execute(select(User).where(User.id == user_id))
    user = result.scalar_one_or_none()

    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Пользователь не найден"
        )

    # Находим файл
    result = await db.execute(select(UploadedFile).where(UploadedFile.id == file_id))
    file_obj = result.scalar_one_or_none()

    if not file_obj:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )

    # Проверяем статус файла
    if file_obj.status != "approved":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Файл не одобрен. Скачивание недоступно."
        )

    # Проверяем существование файла на диске
    file_path = Path(file_obj.file_path)
    if not file_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден на сервере"
        )

    print(f"📥 Скачивание файла: {file_obj.filename} пользователем {user.username}")

    # Возвращаем файл для скачивания
    return FastAPIFileResponse(
        path=str(file_path),
        filename=file_obj.filename,
        media_type='application/octet-stream'
    )





@router.get("/{file_id}/preview")
async def preview_file(
        file_id: int,
        user_id: int,
        db: AsyncSession = Depends(get_db)
):
    """Предпросмотр Word документа"""

    # Проверяем пользователя
    result = await db.execute(select(User).where(User.id == user_id))
    user = result.scalar_one_or_none()

    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Пользователь не найден"
        )

    # Находим файл
    result = await db.execute(select(UploadedFile).where(UploadedFile.id == file_id))
    file_obj = result.scalar_one_or_none()

    if not file_obj:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден"
        )

    # Проверяем статус файла
    if file_obj.status != "approved":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Файл не одобрен. Предпросмотр недоступен."
        )

    # Проверяем тип файла
    if file_obj.file_type not in ["doc", "docx"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Предпросмотр доступен только для Word документов"
        )

    # Проверяем существование файла
    file_path = Path(file_obj.file_path)
    if not file_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл не найден на сервере"
        )

    try:
        # Читаем Word документ
        doc = Document(str(file_path))

        # Извлекаем текст из параграфов
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal"
                })

        # Извлекаем текст из таблиц
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        print(f"👁️ Предпросмотр файла: {file_obj.filename} пользователем {user.username}")

        return {
            "filename": file_obj.filename,
            "paragraphs": paragraphs,
            "tables": tables,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables)
        }

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка чтения документа: {str(e)}"
        )